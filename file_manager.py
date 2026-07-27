import os
import shutil
from werkzeug.utils import secure_filename
import pypdf
import docx
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
WORKSPACES_DIR = os.path.join(BASE_DIR, "NIKI_CORE", "workspaces")

def sanitize_ws_name(name):
    if not name:
        return "general"
    return name.strip().lower().replace(" ", "_")

def sanitize_subfolder(subfolder_path):
    if not subfolder_path:
        return ""
    clean_parts = [secure_filename(part) for part in subfolder_path.replace('\\', '/').split('/') if part and part != '..']
    return "/".join(clean_parts)

def get_target_dir(ws_name, subfolder=""):
    clean_ws = sanitize_ws_name(ws_name)
    clean_sub = sanitize_subfolder(subfolder)
    path = os.path.join(WORKSPACES_DIR, clean_ws, "library", clean_sub)
    os.makedirs(path, exist_ok=True)
    return path

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    extracted_text = ""
    try:
        if ext == ".pdf":
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                extracted_text += (page.extract_text() or "") + "\n"
        elif ext in [".docx", ".doc"]:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        elif ext in [".txt", ".json", ".md"]:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    extracted_text = f.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="cp1251") as f:
                    extracted_text = f.read()
    except Exception as e:
        print(f"Грешка при извличане на текст от {file_path}: {e}")
    return extracted_text.strip()

def save_text_as_docx(ws_name, filename, title, content, subfolder=""):
    library_path = get_target_dir(ws_name, subfolder)
    file_path = os.path.join(library_path, secure_filename(filename))
    doc = Document()
    doc.add_heading(title, level=1)
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    doc.save(file_path)
    return file_path
