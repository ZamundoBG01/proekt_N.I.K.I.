import io
import os
from flask import Blueprint, request, send_file

export_bp = Blueprint('export_bp', __name__)

@export_bp.route('/download_docx_new', methods=['GET', 'POST'])
def download_docx_new():
    # Опитваме се да вземем текста отвсякъде (JSON, Form или URL параметър)
    text = ""
    ws = "general"
    
    if request.method == 'POST':
        if request.is_json:
            data = request.get_json() or {}
            text = data.get('text', '')
            ws = data.get('ws', 'general')
        else:
            text = request.form.get('text', '')
            ws = request.form.get('ws', 'general')
    else:
        # Ако случайно мине през GET, го четем от линка, но с ограничение да не гръмне
        text = request.args.get('text', '')
        ws = request.args.get('ws', 'general')
    
    if not text:
        text = "Няма наличен текст за експорт."

    mem = io.BytesIO()
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(f'Документ от проект: {ws.upper()}', level=1)
        for paragraph in text.split("\n"):
            doc.add_paragraph(paragraph)
        doc.save(mem)
        mem.seek(0)
    except Exception as e:
        print(f"Docx error: {e}")
        mem.write(text.encode('utf-8'))
        mem.seek(0)

    filename = f"NIKI_Response_{ws}.docx"
    mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return send_file(
        mem, 
        as_attachment=True, 
        download_name=filename, 
        mimetype=mimetype
    )

@export_bp.route('/download_pdf_new', methods=['GET', 'POST'])
def download_pdf_new():
    text = ""
    ws = "general"
    
    if request.method == 'POST':
        if request.is_json:
            data = request.get_json() or {}
            text = data.get('text', '')
            ws = data.get('ws', 'general')
        else:
            text = request.form.get('text', '')
            ws = request.form.get('ws', 'general')
    else:
        text = request.args.get('text', '')
        ws = request.args.get('ws', 'general')
        
    if not text:
        text = "Няма наличен текст за експорт."

    mem = io.BytesIO()
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        c = canvas.Canvas(mem, pagesize=letter)
        width, height = letter
        
        font_name = "Helvetica"
        windows_font_path = "C:/Windows/Fonts/arial.ttf"
        linux_font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
        
        if os.path.exists(windows_font_path):
            pdfmetrics.registerFont(TTFont('CustomUnicodeFont', windows_font_path))
            font_name = 'CustomUnicodeFont'
        elif os.path.exists(linux_font_path):
            pdfmetrics.registerFont(TTFont('CustomUnicodeFont', linux_font_path))
            font_name = 'CustomUnicodeFont'

        c.setFont(font_name, 11)
        
        text_lines = text.split("\n")
        y = height - 40
        for line in text_lines:
            if y < 40:
                c.showPage()
                c.setFont(font_name, 11)
                y = height - 40
            c.drawString(40, y, str(line)[:90])
            y -= 20
            
        c.save()
        mem.seek(0)
    except Exception as e:
        print(f"PDF error: {e}")
        mem.write(text.encode('latin-1', errors='ignore'))
        mem.seek(0)

    filename = f"NIKI_Response_{ws}.pdf"
    mimetype = "application/pdf"

    return send_file(
        mem, 
        as_attachment=True, 
        download_name=filename, 
        mimetype=mimetype
    )
