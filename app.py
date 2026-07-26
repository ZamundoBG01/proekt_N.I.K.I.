import os
import re
import shutil
from flask import Flask, render_template, request, jsonify, send_from_directory, Response
import threading

# 1. Първо създаваме приложението
app = Flask(__name__)
app.json.ensure_ascii = False

# 2. Веднага след това регистрираме новите пътища за сваляне
from export_routes import export_bp
app.register_blueprint(export_bp)

# 3. Регистрация на системните пътища за бекъп и възстановяване
from backup_routes import backup_bp
app.register_blueprint(backup_bp)

# 4. Продължават останалите импорти
from database import (
    init_db, get_db_connection, get_workspace_facts, 
    add_workspace_fact, save_chat_message, get_chat_history, clear_workspace_data
)
from file_manager import (
    sanitize_ws_name, sanitize_subfolder, get_target_dir, 
    extract_text_from_file, WORKSPACES_DIR
)
from ai_engine import call_ai_engine, auto_run_worker

init_db()

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/workspaces", methods=["GET", "POST"])
def handle_workspaces():
    if not os.path.exists(WORKSPACES_DIR):
        os.makedirs(WORKSPACES_DIR, exist_ok=True)
    
    conn = get_db_connection()
    if request.method == "POST":
        data = request.get_json() or {}
        raw_name = data.get("name", "")
        ws_name = sanitize_ws_name(raw_name)
        if ws_name:
            ws_path = os.path.join(WORKSPACES_DIR, ws_name)
            os.makedirs(os.path.join(ws_path, "facts"), exist_ok=True)
            os.makedirs(os.path.join(ws_path, "library"), exist_ok=True)
            if conn:
                try:
                    with conn.cursor() as cur:
                        cur.execute("INSERT INTO workspaces (name) VALUES (%s) ON CONFLICT (name) DO NOTHING;", (ws_name,))
                        conn.commit()
                except Exception as e: 
                    print(f"WS Save DB Error: {e}")
                finally: 
                    conn.close()
            return jsonify({"status": "success", "workspace": ws_name})
            
    workspaces = ["general"]
    if conn:
        try:
            with conn.cursor() as cur:
                cur.execute("SELECT name FROM workspaces ORDER BY id ASC;")
                rows = cur.fetchall()
                if rows:
                    workspaces = [r[0] for r in rows]
        except Exception as e: 
            print(f"WS Read DB Error: {e}")
        finally: 
            conn.close()
            
    other_workspaces = sorted([w for w in workspaces if w.lower() != "general"])
    return jsonify({"workspaces": ["general"] + other_workspaces})

@app.route("/workspace_data/<path:ws_name>")
def workspace_data(ws_name):
    clean_ws = sanitize_ws_name(ws_name)
    subfolder = request.args.get("subfolder", "")
    facts = get_workspace_facts(clean_ws)
    chat_history = get_chat_history(clean_ws)
    target_dir = get_target_dir(clean_ws, subfolder)
    
    files = []
    folders = []
    if os.path.exists(target_dir) and os.path.isdir(target_dir):
        try:
            for item in os.listdir(target_dir):
                item_path = os.path.join(target_dir, item)
                if os.path.isdir(item_path):
                    folders.append(item)
                elif os.path.isfile(item_path):
                    files.append(item)
        except Exception as e:
            print(f"Грешка при четене на директорията: {e}")
            
    return jsonify({
        "facts": facts,
        "chat_history": chat_history,
        "tasks": [],
        "files": sorted(files),
        "folders": sorted(folders),
        "current_subfolder": sanitize_subfolder(subfolder)
    })

@app.route("/create_folder", methods=["POST"])
def create_folder():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    subfolder = sanitize_subfolder(data.get("subfolder", ""))
    folder_name = data.get("folder_name", "").strip()
    if not folder_name:
        return jsonify({"message": "Невалидно име на папка."}), 400
    
    new_folder_path = os.path.join(get_target_dir(ws_name, subfolder), folder_name)
    os.makedirs(new_folder_path, exist_ok=True)
    return jsonify({"message": f"Папката '{folder_name}' беше създадена успешно."})

@app.route("/delete_folder", methods=["POST"])
def delete_folder():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    subfolder = sanitize_subfolder(data.get("subfolder", ""))
    folder_name = data.get("folder_name", "").strip()
    folder_path = os.path.join(get_target_dir(ws_name, subfolder), folder_name)
    
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        try:
            shutil.rmtree(folder_path)
            return jsonify({"message": f"Папката '{folder_name}' беше изтрита успешно."})
        except Exception as e:
            return jsonify({"message": f"Грешка при изтриване: {str(e)}"}), 500
    return jsonify({"message": "Папката не бе намерена."}), 404

@app.route("/move_file", methods=["POST"])
def move_file():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    filename = data.get("filename", "")
    source_subfolder = sanitize_subfolder(data.get("source_subfolder", ""))
    target_subfolder = sanitize_subfolder(data.get("target_subfolder", ""))
    
    if not filename:
        return jsonify({"message": "Невалидно име на файл."}), 400
        
    src_dir = get_target_dir(ws_name, source_subfolder)
    target_dir = get_target_dir(ws_name, target_subfolder)
    src_path = os.path.join(src_dir, filename)
    target_path = os.path.join(target_dir, filename)
    
    if not os.path.exists(src_path):
        return jsonify({"message": f"Файлът '{filename}' не съществува."}), 404
        
    try:
        shutil.move(src_path, target_path)
        return jsonify({"message": f"Файлът '{filename}' беше преместен успешно."})
    except Exception as e:
        return jsonify({"message": f"Грешка при преместване: {str(e)}"}), 500

@app.route("/move_file_workspace", methods=["POST"])
def move_file_workspace():
    data = request.json or {}
    source_ws = data.get("source_workspace")
    target_ws = data.get("target_workspace")
    filename = data.get("filename")
    subfolder = data.get("subfolder", "")
    
    import shutil
    import os
    
    base_dir = "workspaces" 
    source_path = os.path.join(base_dir, source_ws, subfolder, filename)
    target_dir = os.path.join(base_dir, target_ws, subfolder)
    
    if not os.path.exists(source_path):
        return jsonify({"status": "error", "message": "Изходният файл не е намерен."}), 404
        
    os.makedirs(target_dir, exist_ok=True)
    target_path = os.path.join(target_dir, filename)
    
    try:
        shutil.move(source_path, target_path)
        return jsonify({"status": "success", "message": "Файлът е преместен успешно."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/delete_all_facts", methods=["POST"])
def delete_all_facts():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("ws", "general"))
    conn = get_db_connection()
    if conn:
        try:
            with conn.cursor() as cur:
                cur.execute("DELETE FROM verified_facts WHERE workspace = %s;", (ws_name,))
                conn.commit()
            return jsonify({"status": "success", "message": "Всички факти бяха изтрити."})
        except Exception as e:
            print(f"Delete All Facts Error: {e}")
            return jsonify({"status": "error", "message": str(e)}), 500
        finally:
            conn.close()
    return jsonify({"status": "error", "message": "Няма връзка с базата данни."}), 500

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json() or {}
    message = data.get("message", "").strip()
    active_ws = sanitize_ws_name(data.get("workspace", "general"))
    auto_run = data.get("auto_run", False)
    
    if not message:
        return jsonify({"reply": "Моля, въведете инструкция.", "monologue": None})
        
    save_chat_message(active_ws, "user", message)
    
    if auto_run:
        t = threading.Thread(target=auto_run_worker, args=(active_ws, message, 3))
        t.start()
        reply_msg = "**Автоматичният офлайн цикъл (Auto-Run) беше стартиран!**"
        save_chat_message(active_ws, "niki", reply_msg, "Стартирана офлайн задача.")
        return jsonify({"reply": reply_msg, "monologue": "Auto-Run Engine Active"})
        
    existing_facts = get_workspace_facts(active_ws)
    library_path = os.path.join(WORKSPACES_DIR, active_ws, "library")
    file_list = []
    library_text = ""
    
    if os.path.exists(library_path):
        for root, dirs, files_in_dir in os.walk(library_path):
            for fname in files_in_dir:
                fpath = os.path.join(root, fname)
                rel_path = os.path.relpath(fpath, library_path)
                file_list.append(rel_path)
                extracted = extract_text_from_file(fpath)
                library_text += f"\n--- ФАЙЛ: {rel_path} ---\n" + (extracted if extracted else "[ПРАЗЕН ФАЙЛ]")

    if "изтрий всичко" in message.lower():
        clear_workspace_data(active_ws)
        reply_msg = f"Всички факти и история в проект **{active_ws.upper()}** бяха изчистени."
        save_chat_message(active_ws, "niki", reply_msg)
        return jsonify({"reply": reply_msg, "monologue": "Изчистване на локалната база данни.", "target_workspace": active_ws})

    if "изтрий факт:" in message.lower():
        fact_to_del = message.lower().replace("изтрий факт:", "").strip()
        conn = get_db_connection()
        if conn:
            try:
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM verified_facts WHERE workspace = %s AND LOWER(content) = %s;", (active_ws, fact_to_del))
                    conn.commit()
            except Exception as e:
                print(f"Delete Fact Error: {e}")
            finally:
                conn.close()
        return jsonify({"status": "success"})

    is_save_command = any(kw in message.lower() for kw in ["запиши", "добави факт", "дневник:"])
    if is_save_command:
        clean_text = re.sub(r"^(запиши предното съобщение|запиши следния факт в базата данни|запиши факт|запиши|добави факт|дневник:)\s*:?", "", message, flags=re.IGNORECASE).strip()
        if not clean_text: 
            clean_text = message
        add_workspace_fact(active_ws, clean_text)
        reply = f"Записах следния факт за постоянно в базата данни на **{active_ws.upper()}**:\n\n> \"{clean_text}\""
        monologue = f"Запис в базата данни: '{clean_text}'"
        save_chat_message(active_ws, "niki", reply, monologue)
        return jsonify({"reply": reply, "monologue": monologue, "target_workspace": active_ws})

    ai_result = call_ai_engine(message, existing_facts, file_list, library_text)
    save_chat_message(active_ws, "niki", ai_result["reply"], ai_result["thought"])
    
    return jsonify({
        "reply": ai_result["reply"],
        "monologue": ai_result["thought"],
        "target_workspace": active_ws
    })

@app.route("/upload", methods=["POST"])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"message": "Няма прикачен файл."}), 400
    file = request.files['file']
    ws_name = sanitize_ws_name(request.form.get("workspace", "general"))
    subfolder = sanitize_subfolder(request.form.get("subfolder", ""))
    
    if file.filename == '':
        return jsonify({"message": "Не е избран файл."}), 400
        
    library_path = get_target_dir(ws_name, subfolder)
    save_path = os.path.join(library_path, file.filename)
    file.save(save_path)
    return jsonify({"message": f"Файлът беше качен успешно."})

@app.route("/download/<path:ws_name>/<path:filename>")
def download_file(ws_name, filename):
    clean_ws = sanitize_ws_name(ws_name)
    library_base = os.path.join(WORKSPACES_DIR, clean_ws, "library")
    return send_from_directory(library_base, filename, as_attachment=True)

@app.route("/download_text_file", methods=["POST"])
def download_text_file():
    text = request.form.get("text", "")
    ws = request.form.get("ws", "general")
    file_format = request.form.get("format", "docx")
    
    import io
    from flask import send_file
    
    mem = io.BytesIO()
    
    if file_format == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        
        c = canvas.Canvas(mem, pagesize=letter)
        width, height = letter
        
        font_name = "Helvetica"
        try:
            windows_font_path = "C:/Windows/Fonts/arial.ttf"
            linux_font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
            
            if os.path.exists(windows_font_path):
                pdfmetrics.registerFont(TTFont('CustomUnicodeFont', windows_font_path))
                font_name = 'CustomUnicodeFont'
            elif os.path.exists(linux_font_path):
                pdfmetrics.registerFont(TTFont('CustomUnicodeFont', linux_font_path))
                font_name = 'CustomUnicodeFont'
        except Exception as e:
            print("Font load error:", e)

        c.setFont(font_name, 11)
        
        text_lines = text.split("\n")
        y = height - 40
        for line in text_lines:
            if y < 40:
                c.showPage()
                c.setFont(font_name, 11)
                y = height - 40
            c.drawString(40, y, str(line)[:90])
            y -= 20
            
        c.save()
        mem.seek(0)
        filename = f"NIKI_Response_{ws}.pdf"
        mimetype = "application/pdf"
        
    else:
        from docx import Document
        
        doc = Document()
        doc.add_heading(f'Документ от проект: {ws.upper()}', level=1)
        
        for paragraph in text.split("\n"):
            doc.add_paragraph(paragraph)
            
        doc.save(mem)
        mem.seek(0)
        filename = f"NIKI_Response_{ws}.docx"
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
    return send_file(mem, as_attachment=True, download_name=filename, mimetype=mimetype)
    
@app.route("/delete_file", methods=["POST"])
def delete_file():
    data = request.get_json() or {}
    ws_name = sanitize_ws_name(data.get("workspace", "general"))
    subfolder = sanitize_subfolder(data.get("subfolder", ""))
    filename = data.get("filename", "")
    
    if not filename:
        return jsonify({"message": "Невалидно име на файл."}), 400
        
    target_dir = get_target_dir(ws_name, subfolder)
    file_path = os.path.join(target_dir, filename)
    
    if os.path.exists(file_path) and os.path.isfile(file_path):
        try:
            os.remove(file_path)
            return jsonify({"message": f"Файлът беше изтрит успешно."})
        except Exception as e:
            return jsonify({"message": f"Грешка при изтриване: {str(e)}"}), 500
    return jsonify({"message": "Файлът не бе намерен на диска."}), 404

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
